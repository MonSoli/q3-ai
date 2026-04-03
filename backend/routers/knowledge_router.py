import uuid
import logging
import asyncio
from datetime import datetime, timezone
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query, BackgroundTasks
from pydantic import BaseModel

from auth import get_current_user
from database import get_db
from models import CreateFolderRequest, RenameFolderRequest, MoveDocumentRequest
from config import MAX_UPLOAD_SIZE
from data_shield import anonymize_text, store_vault_entries, deanonymize_document, load_vault_entries

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/knowledge", tags=["knowledge"])


class DocumentUpdate(BaseModel):
    content: str


@router.get("/folders")
async def get_folders(
    parent_id: Optional[str] = Query(None),
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    if parent_id:
        cursor = await db.execute(
            "SELECT id, name, parent_id, created_at, updated_at FROM knowledge_folders WHERE parent_id = ? ORDER BY name",
            (parent_id,)
        )
    else:
        cursor = await db.execute(
            "SELECT id, name, parent_id, created_at, updated_at FROM knowledge_folders WHERE parent_id IS NULL ORDER BY name"
        )
    rows = await cursor.fetchall()
    return [dict(row) for row in rows]


@router.post("/folders")
async def create_folder(
    req: CreateFolderRequest,
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    if req.parent_id:
        cursor = await db.execute("SELECT id FROM knowledge_folders WHERE id = ?", (req.parent_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Родительская папка не найдена")

    if req.parent_id:
        cursor = await db.execute(
            "SELECT id FROM knowledge_folders WHERE name = ? AND parent_id = ?",
            (req.name, req.parent_id)
        )
    else:
        cursor = await db.execute(
            "SELECT id FROM knowledge_folders WHERE name = ? AND parent_id IS NULL",
            (req.name,)
        )
    if await cursor.fetchone():
        raise HTTPException(status_code=400, detail="Папка с таким именем уже существует")

    folder_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    await db.execute(
        "INSERT INTO knowledge_folders (id, name, parent_id, created_by_id, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?)",
        (folder_id, req.name, req.parent_id, user["id"], now, now)
    )
    await db.commit()

    return {"id": folder_id, "name": req.name, "parent_id": req.parent_id, "message": "Папка создана"}


@router.put("/folders/{folder_id}")
async def rename_folder(
    folder_id: str,
    req: RenameFolderRequest,
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    cursor = await db.execute("SELECT id, parent_id FROM knowledge_folders WHERE id = ?", (folder_id,))
    folder = await cursor.fetchone()
    if not folder:
        raise HTTPException(status_code=404, detail="Папка не найдена")

    parent_id = folder["parent_id"]
    if parent_id:
        cursor = await db.execute(
            "SELECT id FROM knowledge_folders WHERE name = ? AND parent_id = ? AND id != ?",
            (req.name, parent_id, folder_id)
        )
    else:
        cursor = await db.execute(
            "SELECT id FROM knowledge_folders WHERE name = ? AND parent_id IS NULL AND id != ?",
            (req.name, folder_id)
        )
    if await cursor.fetchone():
        raise HTTPException(status_code=400, detail="Папка с таким именем уже существует")

    await db.execute(
        "UPDATE knowledge_folders SET name = ?, updated_at = ? WHERE id = ?",
        (req.name, datetime.now(timezone.utc).isoformat(), folder_id)
    )
    await db.commit()

    return {"message": "Папка переименована"}


@router.delete("/folders/{folder_id}")
async def delete_folder(folder_id: str, user=Depends(get_current_user), db=Depends(get_db)):
    cursor = await db.execute("SELECT id FROM knowledge_folders WHERE id = ?", (folder_id,))
    if not await cursor.fetchone():
        raise HTTPException(status_code=404, detail="Папка не найдена")

    all_ids = [folder_id]
    queue = [folder_id]
    while queue:
        current = queue.pop(0)
        cursor = await db.execute("SELECT id FROM knowledge_folders WHERE parent_id = ?", (current,))
        children = await cursor.fetchall()
        for child in children:
            all_ids.append(child["id"])
            queue.append(child["id"])

    placeholders = ",".join("?" for _ in all_ids)
    await db.execute(f"DELETE FROM knowledge_documents WHERE folder_id IN ({placeholders})", all_ids)

    await db.execute("DELETE FROM knowledge_folders WHERE id = ?", (folder_id,))
    await db.commit()

    return {"message": "Папка удалена"}


@router.get("/documents")
async def get_documents(
    folder_id: Optional[str] = Query(None, alias="folder_id"),
    show_all: bool = Query(False),
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    if show_all:
        cursor = await db.execute(
            """SELECT id, filename, file_type, file_size, folder_id, uploaded_by,
                      doc_type, doc_type_label, is_indexed,
                      created_at, updated_at
               FROM knowledge_documents ORDER BY created_at DESC"""
        )
    elif folder_id == "root" or folder_id is None:
        cursor = await db.execute(
            """SELECT id, filename, file_type, file_size, folder_id, uploaded_by,
                      doc_type, doc_type_label, is_indexed,
                      created_at, updated_at
               FROM knowledge_documents WHERE folder_id IS NULL ORDER BY created_at DESC"""
        )
    else:
        cursor = await db.execute(
            """SELECT id, filename, file_type, file_size, folder_id, uploaded_by,
                      doc_type, doc_type_label, is_indexed,
                      created_at, updated_at
               FROM knowledge_documents WHERE folder_id = ? ORDER BY created_at DESC""",
            (folder_id,)
        )
    rows = await cursor.fetchall()
    return [dict(row) for row in rows]


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    folder_id: Optional[str] = Form(None),
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    if folder_id:
        cursor = await db.execute("SELECT id FROM knowledge_folders WHERE id = ?", (folder_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Папка не найдена")

    allowed_types = [".txt", ".md", ".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"]
    ext = "." + file.filename.split(".")[-1].lower() if "." in file.filename else ""

    if ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат файла. Разрешены: {', '.join(allowed_types)}"
        )

    file_content = await file.read(MAX_UPLOAD_SIZE + 1)
    if len(file_content) > MAX_UPLOAD_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Файл слишком большой. Максимальный размер: {MAX_UPLOAD_SIZE // (1024 * 1024)} МБ"
        )
    file_size = len(file_content)

    content = ""
    if ext in [".txt", ".md"]:
        try:
            content = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                content = file_content.decode("cp1251")
            except Exception:
                content = file_content.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        try:
            import io
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(file_content))
                content = "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    content = "\n".join(page.extract_text() or "" for page in reader.pages)
                except ImportError:
                    content = "[PDF файл - для извлечения текста установите pypdf]"
        except Exception as e:
            content = f"[Ошибка чтения PDF: {str(e)}]"
    elif ext in [".doc", ".docx"]:
        try:
            import io
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_content))
                content = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                content = "[Word файл - для извлечения текста установите python-docx]"
        except Exception as e:
            content = f"[Ошибка чтения документа: {str(e)}]"
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"]:
        try:
            from ocr_engine import ocr_image_bytes, is_image_file
            content = await ocr_image_bytes(file_content, file.filename)
        except Exception as e:
            content = f"[Ошибка OCR: {str(e)}]"

    doc_type = None
    doc_type_label = None
    try:
        from analytics_engine import classify_document_fast
        classification = classify_document_fast(content, file.filename)
        doc_type = classification.get("type")
        doc_type_label = classification.get("type_label")
    except Exception:
        pass

    # --- Data Shield: обезличивание ---
    original_content = content
    anonymized_content, vault_entries = anonymize_text(content)
    shielded = len(vault_entries) > 0
    if shielded:
        logger.info("Data Shield: document '%s' — %d sensitive values anonymized", file.filename, len(vault_entries))

    doc_id = str(uuid.uuid4())
    uploaded_by_name = f"{user.get('last_name', '')} {user.get('first_name', '')}".strip() or user.get('email', '')

    now = datetime.now(timezone.utc).isoformat()
    await db.execute(
        """INSERT INTO knowledge_documents
           (id, filename, file_type, file_size, content, folder_id, uploaded_by_id, uploaded_by,
            doc_type, doc_type_label, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (
            doc_id, file.filename, ext, file_size, anonymized_content,
            folder_id, user["id"], uploaded_by_name,
            doc_type, doc_type_label,
            now, now,
        )
    )

    # Сохраняем зашифрованные значения в vault
    if vault_entries:
        await store_vault_entries(db, doc_id, vault_entries)

    await db.commit()

    if content and not content.startswith("["):
        async def _auto_index():
            try:
                from rag_engine import index_document
                await index_document(doc_id, file.filename, anonymized_content, folder_id)
                logger.info("Auto-indexed document: %s", file.filename)
            except Exception as e:
                logger.warning("Auto-index failed for %s: %s", file.filename, e)
        asyncio.create_task(_auto_index())

    return {
        "id": doc_id,
        "filename": file.filename,
        "file_type": ext,
        "file_size": file_size,
        "folder_id": folder_id,
        "doc_type": doc_type,
        "doc_type_label": doc_type_label,
        "data_shielded": shielded,
        "shielded_count": len(vault_entries),
        "message": "Документ успешно загружен" + (f" (обезличено: {len(vault_entries)} значений)" if shielded else "")
    }


@router.get("/documents/{doc_id}")
async def get_document(doc_id: str, user=Depends(get_current_user), db=Depends(get_db)):
    cursor = await db.execute(
        """SELECT id, filename, file_type, file_size, content, folder_id, uploaded_by, created_at, updated_at
           FROM knowledge_documents WHERE id = ?""",
        (doc_id,)
    )
    row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Документ не найден")
    result = dict(row)
    # Data Shield: деанонимизация для авторизованного пользователя
    if result.get("content"):
        result["content"] = await deanonymize_document(db, doc_id, result["content"])
    return result


@router.put("/documents/{doc_id}")
async def update_document(
    doc_id: str,
    data: DocumentUpdate,
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    cursor = await db.execute("SELECT id FROM knowledge_documents WHERE id = ?", (doc_id,))
    if not await cursor.fetchone():
        raise HTTPException(status_code=404, detail="Документ не найден")

    # Data Shield: обезличиваем новый контент
    anonymized_content, vault_entries = anonymize_text(data.content)
    if vault_entries:
        await store_vault_entries(db, doc_id, vault_entries)

    new_size = len(anonymized_content.encode("utf-8"))
    await db.execute(
        "UPDATE knowledge_documents SET content = ?, file_size = ?, updated_at = ? WHERE id = ?",
        (anonymized_content, new_size, datetime.now(timezone.utc).isoformat(), doc_id)
    )
    await db.commit()
    return {"message": "Документ обновлён"}


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, user=Depends(get_current_user), db=Depends(get_db)):
    cursor = await db.execute("SELECT id FROM knowledge_documents WHERE id = ?", (doc_id,))
    if not await cursor.fetchone():
        raise HTTPException(status_code=404, detail="Документ не найден")

    await db.execute("DELETE FROM knowledge_documents WHERE id = ?", (doc_id,))
    await db.commit()
    return {"message": "Документ удалён"}


@router.post("/documents/{doc_id}/copy")
async def copy_document(doc_id: str, user=Depends(get_current_user), db=Depends(get_db)):
    cursor = await db.execute(
        "SELECT filename, file_type, file_size, content, folder_id FROM knowledge_documents WHERE id = ?",
        (doc_id,)
    )
    row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Документ не найден")

    doc = dict(row)
    new_id = str(uuid.uuid4())
    new_filename = f"Копия - {doc['filename']}"
    uploaded_by_name = f"{user.get('last_name', '')} {user.get('first_name', '')}".strip() or user.get('email', '')

    await db.execute(
        """INSERT INTO knowledge_documents
           (id, filename, file_type, file_size, content, folder_id, uploaded_by_id, uploaded_by, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (
            new_id, new_filename, doc["file_type"], doc["file_size"], doc["content"],
            doc["folder_id"], user["id"], uploaded_by_name,
            datetime.now(timezone.utc).isoformat(), datetime.now(timezone.utc).isoformat(),
        )
    )

    # Data Shield: копируем vault-записи для нового документа
    vault_entries = await load_vault_entries(db, doc_id)
    if vault_entries:
        await store_vault_entries(db, new_id, vault_entries)

    await db.commit()
    return {"id": new_id, "filename": new_filename, "message": "Копия создана"}


@router.post("/documents/{doc_id}/move")
async def move_document(
    doc_id: str,
    req: MoveDocumentRequest,
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    cursor = await db.execute("SELECT id FROM knowledge_documents WHERE id = ?", (doc_id,))
    if not await cursor.fetchone():
        raise HTTPException(status_code=404, detail="Документ не найден")

    if req.folder_id:
        cursor = await db.execute("SELECT id FROM knowledge_folders WHERE id = ?", (req.folder_id,))
        if not await cursor.fetchone():
            raise HTTPException(status_code=404, detail="Папка не найдена")

    await db.execute(
        "UPDATE knowledge_documents SET folder_id = ?, updated_at = ? WHERE id = ?",
        (req.folder_id, datetime.now(timezone.utc).isoformat(), doc_id)
    )
    await db.commit()
    return {"message": "Документ перемещён"}


@router.get("/context")
async def get_knowledge_context(user=Depends(get_current_user), db=Depends(get_db)):
    folder_cursor = await db.execute("SELECT id, name, parent_id FROM knowledge_folders")
    folders = {row["id"]: dict(row) for row in await folder_cursor.fetchall()}

    def get_path(folder_id):
        parts = []
        current = folder_id
        while current and current in folders:
            parts.insert(0, folders[current]["name"])
            current = folders[current]["parent_id"]
        return "/" + "/".join(parts) if parts else "/"

    cursor = await db.execute(
        "SELECT id, filename, content, folder_id FROM knowledge_documents ORDER BY created_at"
    )
    rows = await cursor.fetchall()

    if not rows:
        return {"context": ""}

    context_parts = []
    for row in rows:
        doc = dict(row)
        if doc["content"]:
            # Data Shield: деанонимизация для авторизованного пользователя
            content = await deanonymize_document(db, doc["id"], doc["content"])
            path = get_path(doc["folder_id"])
            full_path = f"{path}/{doc['filename']}" if path != "/" else f"/{doc['filename']}"
            context_parts.append(f"[{full_path}]\n{content}\n")

    return {"context": "\n".join(context_parts)}


def _extract_file_content(file_content: bytes, ext: str) -> str:
    content = ""
    if ext in [".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm",
               ".yaml", ".yml", ".sql", ".py", ".js", ".log", ".ini",
               ".cfg", ".conf", ".sh", ".bat", ".tex", ".svg"]:
        try:
            content = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                content = file_content.decode("cp1251")
            except Exception:
                content = file_content.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        try:
            import io
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(file_content))
                content = "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    content = "\n".join(page.extract_text() or "" for page in reader.pages)
                except ImportError:
                    content = "[PDF — для извлечения текста установите pypdf]"
        except Exception as e:
            content = f"[Ошибка чтения PDF: {str(e)}]"
    elif ext in [".doc", ".docx"]:
        try:
            import io
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_content))
                content = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                content = "[Word — для извлечения текста установите python-docx]"
        except Exception as e:
            content = f"[Ошибка чтения документа: {str(e)}]"
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"]:
        try:
            from ocr_engine import ocr_image_bytes
            import asyncio
            loop = asyncio.get_event_loop()
            content = loop.run_until_complete(ocr_image_bytes(file_content, "image" + ext))
        except Exception as e:
            content = f"[Ошибка OCR: {str(e)}]"
    return content


@router.post("/sort-upload")
async def sort_and_upload(
    files: List[UploadFile] = File(...),
    user=Depends(get_current_user),
    db=Depends(get_db)
):
    from analytics_engine import classify_document_fast, DOCUMENT_TYPES

    if not files:
        raise HTTPException(status_code=400, detail="Файлы не прикреплены")

    now = datetime.now(timezone.utc)
    now_str = now.isoformat()
    date_label = now.strftime("%d.%m.%Y %H:%M")
    uploaded_by_name = f"{user.get('last_name', '')} {user.get('first_name', '')}".strip() or user.get('email', '')

    parent_folder_name = f"Сортировка {date_label}"
    parent_folder_id = str(uuid.uuid4())
    await db.execute(
        "INSERT INTO knowledge_folders (id, name, parent_id, created_by_id, created_at, updated_at) VALUES (?, ?, NULL, ?, ?, ?)",
        (parent_folder_id, parent_folder_name, user["id"], now_str, now_str)
    )

    category_folders = {}
    results = []
    errors = []

    for file in files:
        filename = file.filename or "unknown"
        ext = ("." + filename.split(".")[-1].lower()) if "." in filename else ""

        try:
            file_content = await file.read(MAX_UPLOAD_SIZE + 1)
            if len(file_content) > MAX_UPLOAD_SIZE:
                errors.append({"filename": filename, "error": "Файл слишком большой"})
                continue
            file_size = len(file_content)
        except Exception as e:
            errors.append({"filename": filename, "error": str(e)})
            continue

        content = _extract_file_content(file_content, ext)

        classification = classify_document_fast(content, filename)
        doc_type = classification.get("type", "other")
        doc_type_label = classification.get("type_label", DOCUMENT_TYPES.get(doc_type, "Другое"))

        if doc_type_label not in category_folders:
            folder_id = str(uuid.uuid4())
            await db.execute(
                "INSERT INTO knowledge_folders (id, name, parent_id, created_by_id, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?)",
                (folder_id, doc_type_label, parent_folder_id, user["id"], now_str, now_str)
            )
            category_folders[doc_type_label] = folder_id
        else:
            folder_id = category_folders[doc_type_label]

        # Data Shield: обезличивание
        anonymized_content, vault_entries = anonymize_text(content)

        doc_id = str(uuid.uuid4())
        await db.execute(
            """INSERT INTO knowledge_documents
               (id, filename, file_type, file_size, content, folder_id, uploaded_by_id, uploaded_by,
                doc_type, doc_type_label, created_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                doc_id, filename, ext, file_size, anonymized_content,
                folder_id, user["id"], uploaded_by_name,
                doc_type, doc_type_label,
                now_str, now_str,
            )
        )

        if vault_entries:
            await store_vault_entries(db, doc_id, vault_entries)

        results.append({
            "doc_id": doc_id,
            "filename": filename,
            "category": doc_type_label,
            "doc_type": doc_type,
            "confidence": classification.get("confidence", 0),
            "shielded_count": len(vault_entries),
        })

        if content and not content.startswith("["):
            async def _auto_index(did=doc_id, fn=filename, cnt=anonymized_content, fid=folder_id):
                try:
                    from rag_engine import index_document
                    await index_document(did, fn, cnt, fid)
                except Exception as e:
                    logger.warning("Автоиндексация не удалась для %s: %s", fn, e)
            asyncio.create_task(_auto_index())

    await db.commit()

    categories_summary = {}
    for r in results:
        cat = r["category"]
        if cat not in categories_summary:
            categories_summary[cat] = []
        categories_summary[cat].append(r["filename"])

    return {
        "parent_folder_id": parent_folder_id,
        "parent_folder_name": parent_folder_name,
        "total_files": len(results),
        "total_categories": len(categories_summary),
        "categories": categories_summary,
        "results": results,
        "errors": errors,
    }
